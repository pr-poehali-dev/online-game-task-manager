'''Сборка готовых документов (Excel/Word) по текстовому запросу сотрудника.

Схема работы: модель НЕ генерирует сам файл — она возвращает СТРУКТУРУ документа в JSON
(заголовки, строки таблицы, абзацы), а бинарный .xlsx/.docx собирает уже этот модуль через
openpyxl/python-docx и кладёт в S3. Так надёжнее, чем просить модель отдать base64 файла:
LLM не умеет генерировать корректные zip-контейнеры офисных форматов, зато отлично заполняет
структуру по описанию.
'''

import io
import json
import os
import re
import urllib.error
import urllib.request
import uuid

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from common import (
    _aitunnel_request, _bad, _cors_headers, _current_month, _get_or_create_usage, _ok,
    _service_key, _upload_bytes,
)
from templates import (
    FILL_SYSTEM_PROMPT, MAX_TEMPLATE_BYTES,
    extract_docx_fields, extract_xlsx_fields, fill_docx, fill_xlsx,
)

XLSX_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
DOCX_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

# Модель по умолчанию для разбора запроса в структуру. Задача формальная (заполнить схему JSON),
# рассуждения не нужны — берём быструю и недорогую, иначе каждый счёт-фактура стоил бы как
# полноценный диалог.
DOCUMENT_MODEL = 'gpt-5-mini'

# Ограничения на размер документа — защита от «сделай таблицу на миллион строк»: и модель столько
# не сгенерирует, и функция упрётся в таймаут, и файл будет невозможно открыть.
MAX_ROWS = 500
MAX_COLUMNS = 40
MAX_BLOCKS = 300

# Насколько большую структуру предыдущего документа отдаём модели как основу для правки. Целиком
# слать нельзя: таблица на 500 строк съест весь контекст и подорожает запрос, а для правки вида
# «пересчитай с НДС» модели достаточно увидеть все строки. Ограничиваем по символам.
MAX_SPEC_CONTEXT_CHARS = 60000

EDIT_SYSTEM_PROMPT = (
    'Ты дорабатываешь УЖЕ СОЗДАННЫЙ офисный документ по уточнению пользователя и отвечаешь СТРОГО '
    'одним JSON-объектом той же схемы, что и исходная структура, без markdown-обёртки и пояснений.\n\n'
    'КЛЮЧЕВОЕ ПРАВИЛО: верни ПОЛНУЮ структуру документа целиком — со всеми строками и блоками, '
    'которые были раньше, включая неизменённые. Не возвращай только изменённую часть и не '
    'сокращай данные многоточиями: файл собирается строго из того, что ты вернёшь, поэтому всё '
    'пропущенное будет безвозвратно потеряно.\n'
    'Меняй ТОЛЬКО то, о чём просит пользователь, остальное переноси дословно. Сохраняй kind '
    '(xlsx/docx), порядок колонок и названия — если пользователь явно не просит их изменить. '
    'При пересчёте значений (НДС, скидка, наценка) обнови и итоговые строки.\n'
    # Без этого указания на «добавь ещё три позиции» модель вставляет пустые заглушки вида
    # «Новая позиция 1» с нулями — формально просьба выполнена, но пользоваться таким файлом нельзя.
    'Если пользователь просит добавить строки или разделы, но не уточняет содержание — придумай '
    'ПРАВДОПОДОБНЫЕ данные по теме документа с реалистичными значениями, а не пустые заготовки '
    'вида «Новая позиция 1» с нулями. Отвечай только JSON.'
)

DOCUMENT_SYSTEM_PROMPT = (
    'Ты формируешь СТРУКТУРУ офисного документа по запросу пользователя и отвечаешь СТРОГО одним '
    'JSON-объектом, без markdown-обёртки и пояснений.\n\n'
    'Для таблицы (Excel):\n'
    '{"kind":"xlsx","title":"Название файла","sheets":[{"name":"Лист1",'
    '"columns":["Колонка A","Колонка B"],"rows":[["значение","значение"]],'
    '"totalsRow":["Итого","=СУММА"]}]}\n\n'
    'Для текстового документа (Word):\n'
    '{"kind":"docx","title":"Название файла","blocks":['
    '{"type":"heading","text":"Заголовок","level":1},'
    '{"type":"paragraph","text":"Абзац текста"},'
    '{"type":"bullets","items":["пункт 1","пункт 2"]},'
    '{"type":"table","columns":["A","B"],"rows":[["1","2"]]}]}\n\n'
    'Правила: kind выбирай сам по смыслу запроса (таблица/расчёт/список позиций → xlsx; '
    'письмо, договор, отчёт, инструкция → docx). Все числа в rows передавай числами, а не '
    'строками. Если пользователь не задал конкретные данные — заполни правдоподобным примером по '
    'теме запроса, а не заглушками вида "текст". title — короткое название на языке запроса, без '
    'расширения файла. Отвечай только JSON.'
)


def _strip_json(text: str) -> str:
    '''Достаёт JSON из ответа модели: несмотря на инструкцию, модели регулярно оборачивают ответ
    в ```json-блок или добавляют фразу до/после. Без этой чистки json.loads падал бы на валидном
    по сути ответе.'''
    text = (text or '').strip()
    fence = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
    if fence:
        text = fence.group(1).strip()
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1 and end > start:
        return text[start:end + 1]
    return text


def _safe_filename(title: str, ext: str) -> str:
    '''Имя файла для скачивания: убираем символы, недопустимые в именах файлов Windows/macOS,
    иначе браузер сохранит файл с обрезанным или испорченным именем.'''
    name = re.sub(r'[\\/:*?"<>|\r\n\t]', ' ', str(title or 'Документ')).strip()
    name = re.sub(r'\s+', ' ', name)[:80].strip() or 'Документ'
    return f'{name}.{ext}'


def _cell_value(value):
    '''Приводит значение из JSON к тому, что понимает openpyxl. Словари/списки модель иногда
    подсовывает вместо простого значения — сериализуем, чтобы не потерять данные и не упасть.'''
    if value is None:
        return ''
    if isinstance(value, bool):
        return 'да' if value else 'нет'
    if isinstance(value, (int, float, str)):
        return value
    return json.dumps(value, ensure_ascii=False)


HEADER_FILL = PatternFill('solid', fgColor='4F46E5')
HEADER_FONT = Font(bold=True, color='FFFFFF', size=11)
THIN_BORDER = Border(*[Side(style='thin', color='D4D4D8')] * 4)


def _build_xlsx(spec: dict) -> bytes:
    '''Собирает .xlsx из структуры. Оформление минимальное, но осмысленное: шапка выделена,
    ширина колонок подогнана под содержимое, шапка закреплена при прокрутке — иначе таблица на
    сотню строк нечитаема.'''
    wb = Workbook()
    wb.remove(wb.active)

    sheets = spec.get('sheets') or []
    if not sheets:
        sheets = [{'name': 'Лист1', 'columns': spec.get('columns') or [], 'rows': spec.get('rows') or []}]

    for index, sheet in enumerate(sheets[:10]):
        if not isinstance(sheet, dict):
            continue
        # Имя листа в Excel ограничено 31 символом и не может содержать : \ / ? * [ ]
        raw_name = str(sheet.get('name') or f'Лист{index + 1}')
        name = re.sub(r'[:\\/?*\[\]]', ' ', raw_name)[:31].strip() or f'Лист{index + 1}'
        ws = wb.create_sheet(title=name)

        columns = [str(c) for c in (sheet.get('columns') or [])][:MAX_COLUMNS]
        rows = [r for r in (sheet.get('rows') or []) if isinstance(r, (list, tuple))][:MAX_ROWS]

        if columns:
            ws.append(columns)
            for cell in ws[1]:
                cell.fill = HEADER_FILL
                cell.font = HEADER_FONT
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                cell.border = THIN_BORDER
            ws.freeze_panes = 'A2'

        for row in rows:
            ws.append([_cell_value(v) for v in row[:MAX_COLUMNS]])

        totals = sheet.get('totalsRow')
        if isinstance(totals, (list, tuple)) and totals:
            ws.append([_cell_value(v) for v in totals[:MAX_COLUMNS]])
            for cell in ws[ws.max_row]:
                cell.font = Font(bold=True)

        # Ширина колонок по самому длинному значению — иначе всё слипается в узкие столбцы.
        width_source = ([columns] if columns else []) + [list(r) for r in rows]
        for col_index in range(1, min(len(columns) or MAX_COLUMNS, MAX_COLUMNS) + 1):
            longest = 0
            for row in width_source:
                if col_index - 1 < len(row):
                    longest = max(longest, len(str(_cell_value(row[col_index - 1]))))
            ws.column_dimensions[get_column_letter(col_index)].width = min(max(longest + 3, 10), 60)

        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.border = THIN_BORDER
                cell.alignment = Alignment(vertical='top', wrap_text=True)

    if not wb.sheetnames:
        wb.create_sheet(title='Лист1')

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def _build_docx(spec: dict) -> bytes:
    '''Собирает .docx из списка блоков (заголовки, абзацы, списки, таблицы).'''
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = str(spec.get('title') or '').strip()
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    blocks = spec.get('blocks') or []
    if not blocks and spec.get('text'):
        blocks = [{'type': 'paragraph', 'text': spec['text']}]

    for block in blocks[:MAX_BLOCKS]:
        if not isinstance(block, dict):
            continue
        block_type = str(block.get('type') or 'paragraph')

        if block_type == 'heading':
            level = block.get('level')
            level = level if isinstance(level, int) and 1 <= level <= 4 else 1
            doc.add_heading(str(block.get('text') or ''), level=level)
        elif block_type == 'bullets':
            for item in (block.get('items') or [])[:MAX_ROWS]:
                doc.add_paragraph(str(item), style='List Bullet')
        elif block_type == 'numbered':
            for item in (block.get('items') or [])[:MAX_ROWS]:
                doc.add_paragraph(str(item), style='List Number')
        elif block_type == 'table':
            columns = [str(c) for c in (block.get('columns') or [])][:MAX_COLUMNS]
            rows = [r for r in (block.get('rows') or []) if isinstance(r, (list, tuple))][:MAX_ROWS]
            if not columns and not rows:
                continue
            width = len(columns) or max((len(r) for r in rows), default=1)
            table = doc.add_table(rows=0, cols=width)
            table.style = 'Light Grid Accent 1'
            if columns:
                cells = table.add_row().cells
                for i, column in enumerate(columns[:width]):
                    cells[i].text = column
                    for paragraph in cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for i, value in enumerate(list(row)[:width]):
                    cells[i].text = str(_cell_value(value))
        elif block_type == 'pagebreak':
            doc.add_page_break()
        else:
            doc.add_paragraph(str(block.get('text') or ''))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _fill_template_flow(cur, conn, schema, me, api_key, model, prompt,
                        template_url, template_name, chat_id,
                        user_msg_id, user_created_at, spent, limit_):
    '''Заполняет бланк сотрудника данными от модели, сохраняя оформление исходного файла.

    Порядок: скачиваем бланк из S3 → достаём список плейсхолдеров {{Поле}} → просим модель
    подобрать значения → подставляем их в ИСХОДНЫЙ файл (openpyxl/python-docx открывают его как
    есть, поэтому шрифты, логотип, колонтитулы и формулы остаются нетронутыми) → кладём результат
    в S3 и отдаём ссылку.
    '''
    kind = 'docx' if template_name.lower().endswith('.docx') else 'xlsx' if template_name.lower().endswith('.xlsx') else ''
    if not kind:
        kind = 'docx' if template_url.lower().endswith('.docx') else 'xlsx'

    try:
        with urllib.request.urlopen(template_url, timeout=25) as resp:
            raw_template = resp.read(MAX_TEMPLATE_BYTES + 1)
    except Exception as e:
        print(f'[ai/documents] template download error: {e}')
        cur.close(); conn.close()
        return _bad_message('template_download_failed', 'Не удалось прочитать загруженный бланк — попробуйте загрузить файл заново', chat_id, user_msg_id)

    if len(raw_template) > MAX_TEMPLATE_BYTES:
        cur.close(); conn.close()
        return _bad_message('template_too_large', 'Бланк слишком большой — максимум 8 МБ', chat_id, user_msg_id)

    try:
        fields = extract_docx_fields(raw_template) if kind == 'docx' else extract_xlsx_fields(raw_template)
    except Exception as e:
        print(f'[ai/documents] template parse error: {e}')
        cur.close(); conn.close()
        return _bad_message('template_broken', 'Не удалось открыть бланк — убедитесь, что это файл Word (.docx) или Excel (.xlsx)', chat_id, user_msg_id)

    if not fields:
        cur.close(); conn.close()
        return _bad_message(
            'template_no_fields',
            'В бланке не найдено ни одного поля. Отметьте места для заполнения двойными фигурными скобками, например {{Заказчик}} или {{Сумма}}',
            chat_id, user_msg_id,
        )

    data, err = _aitunnel_request('/chat/completions', api_key, {
        'model': model,
        'messages': [
            {'role': 'system', 'content': FILL_SYSTEM_PROMPT},
            {'role': 'user', 'content': (
                f'Поля бланка: {json.dumps(fields, ensure_ascii=False)}\n\n'
                f'Задание: {prompt}'
            )},
        ],
        'response_format': {'type': 'json_object'},
    }, timeout=90)
    if err:
        cur.close(); conn.close()
        status, payload_err = err
        payload_err['userMessageId'] = user_msg_id
        payload_err['chatId'] = chat_id
        return {'statusCode': status, 'headers': _cors_headers(), 'body': json.dumps(payload_err)}

    choices = data.get('choices') or [{}]
    raw_answer = (choices[0].get('message') or {}).get('content') or ''
    used_model = data.get('model') or model
    cost_rub = (data.get('usage') or {}).get('cost_rub') or 0

    try:
        parsed = json.loads(_strip_json(raw_answer))
        values = parsed.get('values') if isinstance(parsed, dict) else None
        if not isinstance(values, dict):
            raise ValueError('no values')
    except Exception:
        cur.execute(
            f"UPDATE {schema}.ai_usage SET spent_rub = spent_rub + %s WHERE user_id = %s AND month = %s",
            (cost_rub, me['id'], _current_month())
        )
        conn.commit(); cur.close(); conn.close()
        return _bad_message('document_parse_failed', 'Модель вернула данные в неожиданном виде — попробуйте переформулировать запрос', chat_id, user_msg_id, status=502)

    values = {str(k): ('' if v is None else v) for k, v in values.items()}
    try:
        filled = fill_docx(raw_template, values) if kind == 'docx' else fill_xlsx(raw_template, values)
    except Exception as e:
        print(f'[ai/documents] template fill error: {e}')
        cur.close(); conn.close()
        return _bad_message('template_fill_failed', 'Не удалось заполнить бланк — попробуйте ещё раз', chat_id, user_msg_id, status=502)

    content_type = DOCX_CONTENT_TYPE if kind == 'docx' else XLSX_CONTENT_TYPE
    base_name = re.sub(r'\.(docx|xlsx)$', '', template_name, flags=re.I) or 'Документ'
    filename = _safe_filename(f'{base_name} (заполнено)', kind)
    url = _upload_bytes(filled, kind, content_type, 'documents')

    attachment = {
        'id': f'doc-{user_msg_id}', 'name': filename, 'url': url,
        'size': len(filled), 'contentType': content_type,
    }
    filled_count = sum(1 for f in fields if f in values)
    missing = [f for f in fields if f not in values]
    summary = f'Заполнил ваш бланк «{base_name}» — {filled_count} из {len(fields)} полей.'
    if missing:
        summary += ' Не удалось подобрать: ' + ', '.join(missing[:5]) + ('…' if len(missing) > 5 else '')

    cur.execute(
        f"INSERT INTO {schema}.ai_messages (chat_id, role, content, attachments, model, cost_rub) "
        f"VALUES (%s, 'assistant', %s, %s, %s, %s) RETURNING id, created_at",
        (chat_id, summary, json.dumps([attachment]), used_model, cost_rub)
    )
    assistant_msg_id, assistant_created_at = cur.fetchone()
    cur.execute(f"UPDATE {schema}.ai_chats SET updated_at = NOW() WHERE id = %s", (chat_id,))
    cur.execute(
        f"UPDATE {schema}.ai_usage SET spent_rub = spent_rub + %s WHERE user_id = %s AND month = %s",
        (cost_rub, me['id'], _current_month())
    )
    conn.commit(); cur.close(); conn.close()

    return _ok({
        'chatId': chat_id,
        'userMessage': {'id': user_msg_id, 'role': 'user', 'content': prompt, 'createdAt': user_created_at.isoformat()},
        'assistantMessage': {
            'id': assistant_msg_id, 'role': 'assistant', 'content': summary,
            'attachments': [attachment], 'model': used_model, 'costRub': float(cost_rub),
            'createdAt': assistant_created_at.isoformat(),
        },
        'usage': {'spentRub': spent + float(cost_rub), 'limitRub': limit_},
    })


def _bad_message(code, message, chat_id, user_msg_id, status=400):
    '''Ошибка заполнения бланка с понятным текстом для сотрудника — сообщение пользователя уже
    записано в диалог, поэтому возвращаем его id, чтобы фронт корректно снял «отправку».'''
    return {'statusCode': status, 'headers': _cors_headers(), 'body': json.dumps({
        'error': code, 'message': message, 'chatId': chat_id, 'userMessageId': user_msg_id,
    })}


def handle_generate_document(cur, conn, schema, me, body, qs):
    '''Собирает готовый Excel/Word по текстовому запросу сотрудника и возвращает ссылку на файл.

    Модель отдаёт структуру документа в JSON (DOCUMENT_SYSTEM_PROMPT), а сам файл собирается
    здесь через openpyxl/python-docx и кладётся в S3 — модели не умеют отдавать бинарные офисные
    форматы. Формат выбирает модель по смыслу запроса, но его можно задать явно полем format.
    '''
    prompt = (body.get('prompt') or body.get('content') or '').strip()
    chat_id = body.get('chatId')
    model = (body.get('model') or '').strip() or DOCUMENT_MODEL
    # format='xlsx'|'docx' — явный выбор сотрудника в интерфейсе; 'auto' отдаёт решение модели.
    wanted = (body.get('format') or 'auto').strip().lower()
    # baseMessageId — дорабатывать КОНКРЕТНЫЙ документ из переписки (кнопка «Доработать» у нужного
    # файла). Если не передан, но в диалоге уже есть документ — правим последний, т.к. уточнение
    # вида «добавь ещё три позиции» почти всегда относится к только что полученному файлу.
    base_message_id = body.get('baseMessageId')
    # templateUrl — бланк сотрудника (его собственный .docx/.xlsx, загруженный через
    # upload_attachment). Если он передан, документ НЕ собирается с нуля: в готовый файл
    # подставляются значения в плейсхолдеры {{Поле}}, чтобы полностью сохранить оформление.
    template_url = (body.get('templateUrl') or '').strip()
    template_name = (body.get('templateName') or '').strip()
    if not prompt:
        cur.close(); conn.close()
        return _bad('bad_request')

    spent, limit_ = _get_or_create_usage(cur, schema, me['id'])
    if spent >= limit_:
        cur.close(); conn.close()
        return {'statusCode': 403, 'headers': _cors_headers(), 'body': json.dumps({'error': 'limit_exceeded', 'spentRub': spent, 'limitRub': limit_})}

    api_key = _service_key(cur, schema, 'AITUNNEL_API_KEY')
    if not api_key:
        cur.close(); conn.close()
        return _bad('aitunnel_not_configured')

    if chat_id:
        cur.execute(f"SELECT id FROM {schema}.ai_chats WHERE id = %s AND user_id = %s", (chat_id, me['id']))
        if not cur.fetchone():
            cur.close(); conn.close()
            return _bad('not_found', 404)
    else:
        title = 'Документ: ' + (prompt[:45] + ('…' if len(prompt) > 45 else ''))
        cur.execute(
            f"INSERT INTO {schema}.ai_chats (user_id, title, mode, model) VALUES (%s, %s, 'document', %s) RETURNING id",
            (me['id'], title, model)
        )
        chat_id = cur.fetchone()[0]

    # Структура документа, который дорабатываем. Ищем либо конкретное сообщение (baseMessageId),
    # либо последний собранный документ в этом же диалоге. Проверка chat_id/user_id обязательна:
    # иначе по чужому messageId можно было бы вытащить документ другого сотрудника.
    base_spec = None
    if chat_id:
        if base_message_id:
            cur.execute(
                f"SELECT m.doc_spec FROM {schema}.ai_messages m "
                f"JOIN {schema}.ai_chats c ON c.id = m.chat_id "
                f"WHERE m.id = %s AND m.chat_id = %s AND c.user_id = %s AND m.doc_spec IS NOT NULL",
                (base_message_id, chat_id, me['id'])
            )
        else:
            cur.execute(
                f"SELECT m.doc_spec FROM {schema}.ai_messages m "
                f"JOIN {schema}.ai_chats c ON c.id = m.chat_id "
                f"WHERE m.chat_id = %s AND c.user_id = %s AND m.doc_spec IS NOT NULL "
                f"ORDER BY m.id DESC LIMIT 1",
                (chat_id, me['id'])
            )
        row = cur.fetchone()
        if row and row[0]:
            base_spec = row[0]

    cur.execute(
        f"INSERT INTO {schema}.ai_messages (chat_id, role, content) VALUES (%s, 'user', %s) RETURNING id, created_at",
        (chat_id, prompt)
    )
    user_msg_id, user_created_at = cur.fetchone()

    # Ветка заполнения бланка — полностью отдельный путь: файл сотрудника правится точечно,
    # структура документа не строится и в doc_spec не пишется (дорабатывать бланк уточнениями
    # нельзя — нужно заново заполнить исходный бланк).
    if template_url:
        return _fill_template_flow(
            cur, conn, schema, me, api_key, model, prompt,
            template_url, template_name, chat_id,
            user_msg_id, user_created_at, spent, limit_,
        )

    if base_spec is not None:
        # Режим доработки: модель получает предыдущую структуру и возвращает её целиком с учётом
        # правки. Слишком большую структуру не отправляем — вместо неверного результата лучше
        # честно собрать документ заново по описанию.
        base_json = json.dumps(base_spec, ensure_ascii=False)
        if len(base_json) > MAX_SPEC_CONTEXT_CHARS:
            base_spec = None

    if base_spec is not None:
        instruction = EDIT_SYSTEM_PROMPT
        user_content = (
            f'Текущая структура документа:\n{base_json}\n\n'
            f'Уточнение пользователя: {prompt}'
        )
    else:
        instruction = DOCUMENT_SYSTEM_PROMPT
        user_content = prompt
        if wanted in ('xlsx', 'docx'):
            instruction += f'\n\nПользователь явно выбрал формат: используй kind="{wanted}".'

    data, err = _aitunnel_request('/chat/completions', api_key, {
        'model': model,
        'messages': [
            {'role': 'system', 'content': instruction},
            {'role': 'user', 'content': user_content},
        ],
        'response_format': {'type': 'json_object'},
    }, timeout=90)
    if err:
        cur.close(); conn.close()
        status, payload_err = err
        payload_err['userMessageId'] = user_msg_id
        payload_err['chatId'] = chat_id
        return {'statusCode': status, 'headers': _cors_headers(), 'body': json.dumps(payload_err)}

    choices = data.get('choices') or [{}]
    raw_answer = (choices[0].get('message') or {}).get('content') or ''
    used_model = data.get('model') or model
    usage = data.get('usage') or {}
    cost_rub = usage.get('cost_rub') or 0

    try:
        spec = json.loads(_strip_json(raw_answer))
        if not isinstance(spec, dict):
            raise ValueError('not an object')
    except Exception:
        # Модель не вернула валидную структуру — деньги за запрос уже списаны провайдером, поэтому
        # фиксируем расход, но сообщаем об ошибке понятным текстом.
        cur.execute(
            f"UPDATE {schema}.ai_usage SET spent_rub = spent_rub + %s WHERE user_id = %s AND month = %s",
            (cost_rub, me['id'], _current_month())
        )
        conn.commit()
        cur.close(); conn.close()
        return {'statusCode': 502, 'headers': _cors_headers(), 'body': json.dumps({
            'error': 'document_parse_failed',
            'message': 'Модель вернула документ в неожиданном виде. Попробуйте переформулировать запрос или выбрать другую модель.',
            'chatId': chat_id, 'userMessageId': user_msg_id,
        })}

    kind = str(spec.get('kind') or wanted or 'xlsx').lower()
    if wanted in ('xlsx', 'docx'):
        kind = wanted
    # При доработке формат наследуется от исходного документа: уточнение «добавь позиции» не должно
    # неожиданно превратить таблицу Excel в документ Word, даже если модель ошиблась с kind.
    if base_spec is not None and wanted not in ('xlsx', 'docx'):
        base_kind = str(base_spec.get('kind') or '').lower()
        if base_kind in ('xlsx', 'docx'):
            kind = base_kind
    if kind not in ('xlsx', 'docx'):
        kind = 'xlsx' if spec.get('sheets') or spec.get('rows') else 'docx'

    if kind == 'xlsx':
        raw = _build_xlsx(spec)
        content_type = XLSX_CONTENT_TYPE
    else:
        raw = _build_docx(spec)
        content_type = DOCX_CONTENT_TYPE

    doc_title = str(spec.get('title') or '').strip() or prompt[:60]
    filename = _safe_filename(doc_title, kind)
    url = _upload_bytes(raw, kind, content_type, 'documents')

    attachment = {
        'id': f'doc-{user_msg_id}',
        'name': filename,
        'url': url,
        'size': len(raw),
        'contentType': content_type,
    }
    # Короткое человекочитаемое описание вместо сырого JSON — в ленте показывается как текст ответа.
    updated = base_spec is not None
    if kind == 'xlsx':
        sheets = spec.get('sheets') or []
        total_rows = sum(len(s.get('rows') or []) for s in sheets if isinstance(s, dict))
        if updated:
            was_rows = sum(len(s.get('rows') or []) for s in (base_spec.get('sheets') or []) if isinstance(s, dict))
            delta = total_rows - was_rows
            change = f', строк: {was_rows} → {total_rows}' if delta else ''
            summary = f'Обновил таблицу «{doc_title}»{change}.'
        else:
            summary = f'Готова таблица «{doc_title}» — {total_rows} строк.'
    else:
        summary = f'Обновил документ «{doc_title}».' if updated else f'Готов документ «{doc_title}».'

    cur.execute(
        f"INSERT INTO {schema}.ai_messages (chat_id, role, content, attachments, model, cost_rub, doc_spec) "
        f"VALUES (%s, 'assistant', %s, %s, %s, %s, %s) RETURNING id, created_at",
        (chat_id, summary, json.dumps([attachment]), used_model, cost_rub, json.dumps(spec, ensure_ascii=False))
    )
    assistant_msg_id, assistant_created_at = cur.fetchone()
    cur.execute(f"UPDATE {schema}.ai_chats SET updated_at = NOW() WHERE id = %s", (chat_id,))
    cur.execute(
        f"UPDATE {schema}.ai_usage SET spent_rub = spent_rub + %s WHERE user_id = %s AND month = %s",
        (cost_rub, me['id'], _current_month())
    )
    conn.commit()
    cur.close(); conn.close()

    return _ok({
        'chatId': chat_id,
        'userMessage': {'id': user_msg_id, 'role': 'user', 'content': prompt, 'createdAt': user_created_at.isoformat()},
        'assistantMessage': {
            'id': assistant_msg_id, 'role': 'assistant', 'content': summary,
            'attachments': [attachment], 'model': used_model, 'costRub': float(cost_rub),
            'createdAt': assistant_created_at.isoformat(),
            # hasDocSpec — по этому признаку интерфейс показывает у сообщения кнопку «Доработать».
            'hasDocSpec': True, 'documentUpdated': updated,
        },
        'usage': {'spentRub': spent + float(cost_rub), 'limitRub': limit_},
    })


def _multipart_body(fields: dict, file_field: str, filename: str, file_bytes: bytes, content_type: str):
    '''Собирает тело multipart/form-data вручную: в стандартной библиотеке Python готового
    сборщика нет, а тянуть requests в зависимости функции ради одного запроса незачем.'''
    boundary = '----AiDocBoundary' + uuid.uuid4().hex
    parts = []
    for name, value in fields.items():
        parts.append(
            f'--{boundary}\r\nContent-Disposition: form-data; name="{name}"\r\n\r\n{value}\r\n'.encode('utf-8')
        )
    parts.append(
        f'--{boundary}\r\nContent-Disposition: form-data; name="{file_field}"; filename="{filename}"\r\n'
        f'Content-Type: {content_type}\r\n\r\n'.encode('utf-8')
    )
    parts.append(file_bytes)
    parts.append(f'\r\n--{boundary}--\r\n'.encode('utf-8'))
    return b''.join(parts), f'multipart/form-data; boundary={boundary}'


def _tg_send_document(chat_id, file_url, filename, caption):
    '''Отправляет файл в личные сообщения Telegram.

    Файл передаём БАЙТАМИ через multipart, а не ссылкой на CDN. Вариант со ссылкой выглядит проще
    (Telegram скачивает сам), но на практике упирается в таймаут: Telegram тянет файл с нашего CDN
    неопределённо долго, и запрос обрывается по urlopen timeout. Скачать документ из S3 самим и
    отдать готовые байты — предсказуемо быстро, документы весят единицы мегабайт.
    Возвращает (True, None) при успехе либо (False, 'текст ошибки') — вызывающий код показывает
    её сотруднику как есть.'''
    token = os.environ.get('TELEGRAM_BOT_TOKEN', '')
    if not token:
        return False, 'Бот Telegram не настроен — обратитесь к администратору'
    if not chat_id:
        return False, 'У получателя не привязан Telegram'

    try:
        with urllib.request.urlopen(file_url, timeout=25) as resp:
            file_bytes = resp.read()
    except Exception as e:
        print(f'[ai/documents] cdn download error: {e}')
        return False, 'Не удалось получить файл из хранилища — попробуйте ещё раз'

    body, content_type_header = _multipart_body(
        {'chat_id': str(chat_id), 'caption': caption[:1024]},
        'document', filename or 'document', file_bytes,
        'application/octet-stream',
    )
    req = urllib.request.Request(
        f'https://api.telegram.org/bot{token}/sendDocument',
        data=body, headers={'Content-Type': content_type_header}, method='POST',
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            json.loads(resp.read().decode('utf-8'))
        return True, None
    except urllib.error.HTTPError as e:
        raw = e.read().decode('utf-8', 'ignore')
        description = ''
        try:
            description = (json.loads(raw) or {}).get('description', '')
        except Exception:
            description = raw
        print(f'[ai/documents] telegram sendDocument HTTP {e.code}: {raw}')
        # Самая частая причина — сотрудник не открывал диалог с ботом: боту запрещено писать первым.
        if 'bot can\'t initiate conversation' in description or 'chat not found' in description:
            return False, 'Получатель ещё не начал диалог с ботом — попросите его написать боту любое сообщение'
        if 'blocked' in description:
            return False, 'Получатель заблокировал бота в Telegram'
        return False, 'Telegram не принял файл — попробуйте ещё раз'
    except Exception as e:
        print(f'[ai/documents] telegram sendDocument error: {e}')
        return False, 'Не удалось связаться с Telegram — попробуйте ещё раз'


def handle_send_document_telegram(cur, conn, schema, me, body, qs):
    '''Отправляет ранее собранный документ в личные сообщения Telegram — себе или коллеге.

    Файл не пересылается через функцию: Telegram скачивает его сам по CDN-ссылке из вложения
    сообщения. Отправить можно только документ из СВОЕГО диалога (проверка user_id), получатель —
    любой активный сотрудник, видимый в списке команды.
    '''
    message_id = body.get('messageId')
    # recipientId не передан → отправляем самому себе (самый частый сценарий).
    recipient_id = body.get('recipientId') or me['id']
    if not message_id:
        cur.close(); conn.close()
        return _bad('bad_request')

    # Вложение берём из БД, а не из тела запроса: иначе можно было бы подсунуть произвольную
    # ссылку и разослать её команде от имени сервиса.
    cur.execute(
        f"SELECT m.attachments FROM {schema}.ai_messages m "
        f"JOIN {schema}.ai_chats c ON c.id = m.chat_id "
        f"WHERE m.id = %s AND c.user_id = %s",
        (message_id, me['id'])
    )
    row = cur.fetchone()
    if not row or not row[0]:
        cur.close(); conn.close()
        return _bad('not_found', 404)

    attachments = [a for a in row[0] if isinstance(a, dict) and a.get('url')]
    document = next(
        (a for a in attachments if a.get('contentType') in (XLSX_CONTENT_TYPE, DOCX_CONTENT_TYPE)),
        None
    )
    if not document:
        cur.close(); conn.close()
        return _bad('not_found', 404)

    cur.execute(
        f"SELECT telegram_id, first_name FROM {schema}.users "
        f"WHERE id = %s AND is_active = true AND telegram_id > 0",
        (recipient_id,)
    )
    recipient = cur.fetchone()
    if not recipient:
        cur.close(); conn.close()
        return {'statusCode': 400, 'headers': _cors_headers(), 'body': json.dumps({
            'error': 'recipient_unavailable',
            'message': 'У получателя не привязан Telegram или сотрудник неактивен',
        })}

    telegram_id, recipient_name = recipient
    caption = f'{document.get("name", "Документ")}\n\nДокумент из раздела AI'
    if recipient_id != me['id']:
        # Имя отправителя подписываем только при отправке коллеге — чтобы получатель понимал,
        # от кого пришёл файл. В _current_user имени нет, поэтому берём отдельным запросом.
        cur.execute(f"SELECT first_name FROM {schema}.users WHERE id = %s", (me['id'],))
        sender_row = cur.fetchone()
        sender = (sender_row[0] or '').strip() if sender_row else ''
        if sender:
            caption += f' · отправил {sender}'

    ok, error_text = _tg_send_document(telegram_id, document['url'], document.get('name'), caption)
    cur.close(); conn.close()
    if not ok:
        return {'statusCode': 502, 'headers': _cors_headers(), 'body': json.dumps({
            'error': 'telegram_failed', 'message': error_text,
        })}
    return _ok({
        'sent': True,
        'recipientName': recipient_name,
        'toSelf': recipient_id == me['id'],
    })

def handle_document_recipients(cur, conn, schema, me, body, qs):
    '''Список сотрудников, которым можно отправить документ в Telegram.

    Берём тех же людей, что видны в списке команды (auth action=team): активные, не скрытые,
    с привязанным Telegram. Сам сотрудник всегда идёт первым с пометкой isSelf — отправка себе
    самый частый сценарий.
    '''
    cur.execute(
        f"SELECT id, COALESCE(NULLIF(nickname, ''), first_name) AS name, last_name, "
        f"(id = %s) AS is_self "
        f"FROM {schema}.users "
        f"WHERE is_active = true AND is_hidden = false AND show_in_team = true AND telegram_id > 0 "
        f"ORDER BY (id = %s) DESC, first_name ASC",
        (me['id'], me['id'])
    )
    recipients = [
        {
            'id': r[0],
            'name': ' '.join(part for part in (r[1], r[2]) if part) or 'Без имени',
            'isSelf': bool(r[3]),
        }
        for r in cur.fetchall()
    ]
    cur.close(); conn.close()
    return _ok({'recipients': recipients})
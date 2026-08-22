'''Разбор файлов проекта на фрагменты для поиска по содержимому (этап 2 плана AI_PROJECTS_PLAN.md).

Зачем фрагменты: отдать модели все файлы проекта целиком нельзя — они не влезут в контекст и будут
стоить очень дорого. Поэтому каждый файл режется на куски по ~1000 символов, а при вопросе
сотрудника подбираются только те несколько кусков, что относятся к вопросу.

Способ поиска — встроенный полнотекстовый поиск PostgreSQL с русской морфологией ('договоры'
находятся по запросу 'договор'). Расширения pgvector на сервере нет, это предусмотренный планом
фолбэк; колонка embedding в ai_file_chunks заведена заранее на случай его появления.

Обход таймаута функции: большой документ невозможно разобрать за один вызов (лимит ~5 секунд),
поэтому разбор идёт ПОРЦИЯМИ — фронт вызывает index_step в цикле, каждый вызов обрабатывает
ограниченный кусок и запоминает позицию в ai_files.index_offset.'''

import io
import json
import os
import re

from common import _bad, _ok, _s3_client

# Размер одного фрагмента и перекрытие между соседними. Перекрытие нужно, чтобы фраза, попавшая на
# границу двух кусков, всё равно нашлась целиком хотя бы в одном из них.
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150

# Сколько символов текста разбираем за ОДИН вызов функции. Значение подобрано с запасом под таймаут
# ~5 секунд: вставка в БД быстрая, основное время съедает скачивание и разбор файла.
CHARS_PER_STEP = 40000

# Максимум текста, который вообще извлекаем из одного файла — защита от гигантских документов,
# которые раздули бы базу и стоимость поиска.
MAX_TEXT_PER_FILE = 400000

# Русская морфология для полнотекстового поиска. 'russian' есть в любой стандартной сборке
# PostgreSQL, отдельная установка не нужна.
#
# ВАЖНО про lower(): на этом сервере база создана без русской локали, поэтому PostgreSQL НЕ
# приводит заглавные кириллические буквы к строчным сам — 'Договоры' превращается в лексему
# 'Договор' (с большой буквы) и не совпадает с запросом 'договор'. Проверено запросом к БД.
# Поэтому и текст, и запрос всегда пропускаем через lower() ЯВНО — тогда морфология работает
# корректно независимо от локали сервера.
TS_CONFIG = 'russian'

TEXT_EXTENSIONS = {
    'txt', 'md', 'markdown', 'csv', 'json', 'xml', 'yaml', 'yml', 'ini', 'cfg', 'conf', 'log',
    'py', 'js', 'ts', 'tsx', 'jsx', 'java', 'c', 'h', 'cpp', 'cs', 'go', 'rs', 'php', 'rb',
    'sh', 'sql', 'html', 'css', 'scss',
}


def _ext(name: str) -> str:
    return (name.rsplit('.', 1)[-1] if '.' in name else '').lower()


def _extract_text(raw: bytes, name: str, content_type: str):
    '''Извлекает читаемый текст из файла. Возвращает (текст, None) либо (None, причина).
    Причина 'unsupported' означает, что в файле просто нет текста (картинка, видео, архив) — это
    штатная ситуация, а не ошибка: такой файл не участвует в поиске, но остаётся в проекте.'''
    ext = _ext(name)
    ctype = (content_type or '').lower()

    if ext == 'pdf' or ctype == 'application/pdf':
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw))
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or '')
                if sum(len(p) for p in parts) > MAX_TEXT_PER_FILE:
                    break
            text = '\n'.join(parts)
            # PDF из сканов не содержит текстового слоя — распознавание картинок здесь не делаем.
            return (text, None) if text.strip() else (None, 'unsupported')
        except Exception:
            return None, 'failed'

    if ext == 'docx':
        try:
            import docx
            document = docx.Document(io.BytesIO(raw))
            parts = [p.text for p in document.paragraphs if p.text]
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))
            text = '\n'.join(parts)
            return (text, None) if text.strip() else (None, 'unsupported')
        except Exception:
            return None, 'failed'

    if ext in ('xlsx', 'xlsm'):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f'# {sheet.title}')
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append(' | '.join(cells))
                    if sum(len(p) for p in parts) > MAX_TEXT_PER_FILE:
                        break
            wb.close()
            text = '\n'.join(parts)
            return (text, None) if text.strip() else (None, 'unsupported')
        except Exception:
            return None, 'failed'

    if ext in TEXT_EXTENSIONS or ctype.startswith('text/') or ctype in ('application/json', 'application/xml'):
        try:
            text = raw[:MAX_TEXT_PER_FILE * 4].decode('utf-8', errors='replace').replace('\x00', '')
            return (text, None) if text.strip() else (None, 'unsupported')
        except Exception:
            return None, 'failed'

    return None, 'unsupported'


def _split_chunks(text: str, start_index: int):
    '''Режет текст на куски по границам абзацев/предложений, чтобы фрагмент не обрывался на
    середине слова. Возвращает список (индекс, текст).'''
    text = re.sub(r'\n{3,}', '\n\n', text)
    chunks = []
    pos = 0
    index = start_index
    while pos < len(text):
        end = min(pos + CHUNK_SIZE, len(text))
        if end < len(text):
            # Ищем ближайший разумный разрыв назад: абзац, затем конец предложения, затем пробел.
            window = text[pos:end]
            for sep in ('\n\n', '\n', '. ', ' '):
                cut = window.rfind(sep)
                if cut > CHUNK_SIZE // 2:
                    end = pos + cut + len(sep)
                    break
        piece = text[pos:end].strip()
        if piece:
            chunks.append((index, piece))
            index += 1
        if end <= pos:
            break
        pos = max(end - CHUNK_OVERLAP, end) if end >= len(text) else end - CHUNK_OVERLAP
        if pos < 0:
            pos = 0
    return chunks


def _fetch_file_bytes(url: str):
    key = None
    marker = '/bucket/'
    public_url = os.environ.get('S3_PUBLIC_URL', '').rstrip('/')
    if public_url and url.startswith(public_url + '/'):
        key = url[len(public_url) + 1:]
    elif marker in url:
        key = url.split(marker, 1)[1]
    if not key:
        return None
    try:
        obj = _s3_client().get_object(Bucket=os.environ.get('S3_BUCKET', 'files'), Key=key)
        return obj['Body'].read()
    except Exception:
        return None


def handle_index_step(cur, conn, schema, me, body, qs):
    '''Один ШАГ разбора: берёт следующий необработанный файл сотрудника (или указанный в fileId),
    извлекает из него порцию текста и складывает фрагменты в ai_file_chunks. Возвращает прогресс,
    чтобы интерфейс мог показать "Обрабатывается…" и вызвать следующий шаг.'''
    file_id = body.get('fileId') or qs.get('fileId')
    project_id = body.get('projectId') or qs.get('projectId')

    if file_id:
        cur.execute(
            f"SELECT id, name, url, content_type, index_status, index_offset, chunks_count, project_id "
            f"FROM {schema}.ai_files WHERE id = %s AND user_id = %s",
            (file_id, me['id'])
        )
    else:
        # Берём файлы только из проектов: файлы вне проектов в поиске не участвуют, разбирать их
        # незачем (лишние деньги и место).
        where_project = "AND project_id = %s" if project_id else "AND project_id IS NOT NULL"
        params = [me['id']] + ([project_id] if project_id else [])
        cur.execute(
            f"SELECT id, name, url, content_type, index_status, index_offset, chunks_count, project_id "
            f"FROM {schema}.ai_files WHERE user_id = %s {where_project} "
            f"AND index_status IN ('pending', 'indexing') ORDER BY id ASC LIMIT 1",
            tuple(params)
        )

    row = cur.fetchone()
    if not row:
        cur.close(); conn.close()
        return _ok({'done': True, 'pending': 0})

    fid, name, url, content_type, status, offset, chunks_count, file_project = row
    if status in ('ready', 'unsupported', 'failed') and not body.get('force'):
        cur.close(); conn.close()
        return _ok({'done': True, 'fileId': fid, 'status': status})

    raw = _fetch_file_bytes(url)
    if raw is None:
        cur.execute(
            f"UPDATE {schema}.ai_files SET index_status = 'failed', index_error = %s WHERE id = %s",
            ('Файл недоступен в хранилище', fid)
        )
        cur.close(); conn.close()
        return _ok({'done': False, 'fileId': fid, 'status': 'failed'})

    text, reason = _extract_text(raw, name, content_type)
    if text is None:
        cur.execute(
            f"UPDATE {schema}.ai_files SET index_status = %s, index_error = %s WHERE id = %s",
            (reason, 'Из этого файла нельзя извлечь текст' if reason == 'unsupported' else 'Не удалось разобрать файл', fid)
        )
        cur.close(); conn.close()
        return _ok({'done': False, 'fileId': fid, 'status': reason})

    text = text[:MAX_TEXT_PER_FILE]
    # Повторный разбор (force) начинается с нуля — старые фрагменты убираем, чтобы не задвоились.
    if body.get('force'):
        cur.execute(f"DELETE FROM {schema}.ai_file_chunks WHERE file_id = %s", (fid,))
        offset, chunks_count = 0, 0

    portion = text[offset:offset + CHARS_PER_STEP]
    chunks = _split_chunks(portion, chunks_count)
    for index, content in chunks:
        cur.execute(
            f"INSERT INTO {schema}.ai_file_chunks (file_id, project_id, user_id, chunk_index, content, tsv) "
            f"VALUES (%s, %s, %s, %s, %s, to_tsvector('{TS_CONFIG}', lower(%s)))",
            (fid, file_project, me['id'], index, content, content)
        )

    new_offset = offset + len(portion)
    finished = new_offset >= len(text)
    cur.execute(
        f"UPDATE {schema}.ai_files SET index_status = %s, index_offset = %s, chunks_count = %s, index_error = '' "
        f"WHERE id = %s",
        ('ready' if finished else 'indexing', new_offset, chunks_count + len(chunks), fid)
    )

    cur.execute(
        f"SELECT COUNT(*) FROM {schema}.ai_files WHERE user_id = %s AND project_id IS NOT NULL "
        f"AND index_status IN ('pending', 'indexing')",
        (me['id'],)
    )
    pending = int(cur.fetchone()[0])
    cur.close(); conn.close()
    return _ok({
        'done': finished and pending == 0,
        'fileId': fid, 'fileName': name,
        'status': 'ready' if finished else 'indexing',
        'progress': round(min(1.0, new_offset / max(1, len(text))), 2),
        'pending': pending,
    })


def search_chunks(cur, schema, user_id, project_id, query, limit=8):
    '''Поиск фрагментов проекта по запросу. Используется и напрямую (кнопка поиска в проекте), и
    ассистентом на следующем этапе. Возвращает фрагменты с именем файла-источника.'''
    query = (query or '').strip()
    if len(query) < 2:
        return []
    cur.execute(
        f"SELECT c.id, c.file_id, f.name, f.url, c.chunk_index, c.content, "
        f"ts_rank(c.tsv, plainto_tsquery('{TS_CONFIG}', lower(%s))) AS rank "
        f"FROM {schema}.ai_file_chunks c JOIN {schema}.ai_files f ON f.id = c.file_id "
        f"WHERE c.user_id = %s AND c.project_id = %s AND c.tsv @@ plainto_tsquery('{TS_CONFIG}', lower(%s)) "
        f"ORDER BY rank DESC LIMIT %s",
        (query, user_id, project_id, query, limit)
    )
    rows = cur.fetchall()
    if not rows:
        # Фолбэк на подстроку: короткие запросы, латиница, артикулы и коды морфология не ловит.
        pattern = '%' + query.replace('\\', '\\\\').replace('%', '\\%').replace('_', '\\_') + '%'
        cur.execute(
            f"SELECT c.id, c.file_id, f.name, f.url, c.chunk_index, c.content, 0 AS rank "
            f"FROM {schema}.ai_file_chunks c JOIN {schema}.ai_files f ON f.id = c.file_id "
            f"WHERE c.user_id = %s AND c.project_id = %s AND c.content ILIKE %s LIMIT %s",
            (user_id, project_id, pattern, limit)
        )
        rows = cur.fetchall()
    return [{
        'chunkId': r[0], 'fileId': r[1], 'fileName': r[2], 'fileUrl': r[3],
        'chunkIndex': r[4], 'content': r[5], 'rank': float(r[6] or 0),
    } for r in rows]


def handle_search_project(cur, conn, schema, me, body, qs):
    '''Поиск по файлам проекта — сотрудник может искать сам, не спрашивая ассистента.'''
    project_id = body.get('projectId') or qs.get('projectId')
    query = body.get('query') or qs.get('query') or ''
    if not project_id:
        cur.close(); conn.close()
        return _bad('bad_request')
    cur.execute(f"SELECT id FROM {schema}.ai_projects WHERE id = %s AND user_id = %s", (project_id, me['id']))
    if not cur.fetchone():
        cur.close(); conn.close()
        return _bad('not_found', 404)
    results = search_chunks(cur, schema, me['id'], project_id, query)
    cur.close(); conn.close()
    return _ok({'results': results})


def handle_index_status(cur, conn, schema, me, body, qs):
    '''Сколько файлов проекта ещё не разобрано — интерфейс по этому числу решает, показывать ли
    индикатор обработки и продолжать ли вызывать index_step.'''
    project_id = body.get('projectId') or qs.get('projectId')
    where = "user_id = %s AND project_id IS NOT NULL"
    params = [me['id']]
    if project_id:
        where = "user_id = %s AND project_id = %s"
        params.append(project_id)
    cur.execute(
        f"SELECT index_status, COUNT(*) FROM {schema}.ai_files WHERE {where} GROUP BY index_status",
        tuple(params)
    )
    counts = {status: int(count) for status, count in cur.fetchall()}
    cur.close(); conn.close()
    return _ok({
        'pending': counts.get('pending', 0) + counts.get('indexing', 0),
        'ready': counts.get('ready', 0),
        'unsupported': counts.get('unsupported', 0),
        'failed': counts.get('failed', 0),
    })

'''Заполнение ПОЛЬЗОВАТЕЛЬСКИХ бланков Word/Excel данными от модели.

Отличие от documents.py: там файл собирается с нуля, здесь — берётся готовый бланк сотрудника и
правится ТОЧЕЧНО, чтобы полностью сохранить его оформление (шрифты, логотип, колонтитулы, печати,
рамки, ширину колонок). Поэтому документ никогда не пересобирается: openpyxl/python-docx
открывают исходный файл и меняют только текст плейсхолдеров.

Плейсхолдеры — {{Поле}} в любом месте бланка (абзац, таблица, колонтитул, ячейка). Модель получает
их список и возвращает значения, подстановка выполняется здесь.
'''

import io
import json
import re

from openpyxl import load_workbook
from docx import Document

# Плейсхолдер: {{Заказчик}}, {{ Дата договора }}, {{СУММА}}. Пробелы вокруг имени игнорируем —
# в реальных бланках их ставят как попало.
PLACEHOLDER_RE = re.compile(r'\{\{\s*([^{}]+?)\s*\}\}')

# Ограничение на размер бланка: файл читается целиком в память облачной функции (256 МБ),
# а осмысленные бланки договоров/смет весят единицы мегабайт.
MAX_TEMPLATE_BYTES = 8 * 1024 * 1024

FILL_SYSTEM_PROMPT = (
    'Ты заполняешь бланк документа. Тебе дают СПИСОК ПОЛЕЙ бланка и задание пользователя. '
    'Ответь СТРОГО одним JSON-объектом вида {"values": {"ИмяПоля": "значение", ...}}, без '
    'markdown-обёртки и пояснений.\n\n'
    'Правила:\n'
    '- Ключи в values — РОВНО те имена полей, что тебе дали, посимвольно, ничего не придумывай '
    'и не переименовывай.\n'
    '- Заполни все поля. Если данных для поля нет в задании — подбери правдоподобное значение '
    'по смыслу названия поля и теме документа (например для «Дата» — конкретную дату, для '
    '«Сумма» — число), но НЕ оставляй заглушки вида «—», «нет данных», «XXX».\n'
    '- Значения — обычный текст без кавычек-ёлочек по краям, без markdown, без переносов строк, '
    'если поле явно не многострочное.\n'
    '- Даты пиши в формате ДД.ММ.ГГГГ, если в бланке рядом не указан другой формат.\n'
    # Excel складывает только настоящие числа: значение «60 000 руб.» попадёт в ячейку текстом,
    # и формула ИТОГО в бланке даст ноль. Поэтому числовые поля просим отдавать голым числом —
    # единицы измерения и так подписаны в шапке бланка.
    '- Если поле числовое (сумма, цена, количество, процент) — верни ТОЛЬКО число без пробелов, '
    'без валюты и без единиц измерения: 60000, а не «60 000 руб.» и не «5 часов». Иначе формулы '
    'в бланке перестанут считать. Единицы измерения указывай только если поле явно текстовое '
    '(например «Срок»).'
)


def _iter_docx_paragraphs(document):
    '''Все абзацы документа Word, включая те, что внутри таблиц и колонтитулов. Обычный
    document.paragraphs их НЕ возвращает, и плейсхолдеры в шапке бланка (там обычно реквизиты
    и номер договора) остались бы незаполненными.'''
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                # Вложенные таблицы встречаются в бланках со сложной вёрсткой.
                for nested in cell.tables:
                    for nested_row in nested.rows:
                        for nested_cell in nested_row.cells:
                            for paragraph in nested_cell.paragraphs:
                                yield paragraph
    for section in document.sections:
        for part in (section.header, section.footer, section.first_page_header,
                     section.first_page_footer, section.even_page_header, section.even_page_footer):
            if part is None:
                continue
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def _replace_in_paragraph(paragraph, values: dict) -> int:
    '''Подставляет значения в абзац, СОХРАНЯЯ форматирование.

    Тонкость Word: один плейсхолдер часто разбит на несколько runs («{{», «Заказ», «чик}}») —
    так бывает после правок или проверки орфографии. Поэтому склеиваем текст абзаца целиком,
    выполняем замену, а результат кладём в ПЕРВЫЙ run (он хранит шрифт/размер/начертание этого
    места), остальные очищаем. Так стиль абзаца не теряется.
    '''
    runs = paragraph.runs
    if not runs:
        return 0
    full_text = ''.join(run.text for run in runs)
    if '{{' not in full_text:
        return 0

    replaced = 0

    def substitute(match):
        nonlocal replaced
        name = match.group(1).strip()
        if name in values:
            replaced += 1
            return str(values[name])
        return match.group(0)

    new_text = PLACEHOLDER_RE.sub(substitute, full_text)
    if not replaced:
        return 0
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ''
    return replaced


def extract_docx_fields(raw: bytes) -> list:
    '''Список имён плейсхолдеров в бланке Word (в порядке появления, без повторов).'''
    document = Document(io.BytesIO(raw))
    fields, seen = [], set()
    for paragraph in _iter_docx_paragraphs(document):
        for name in PLACEHOLDER_RE.findall(''.join(run.text for run in paragraph.runs) or paragraph.text):
            key = name.strip()
            if key and key not in seen:
                seen.add(key)
                fields.append(key)
    return fields


def fill_docx(raw: bytes, values: dict) -> bytes:
    '''Заполняет бланк Word. Файл открывается как есть — всё оформление сохраняется.'''
    document = Document(io.BytesIO(raw))
    for paragraph in _iter_docx_paragraphs(document):
        _replace_in_paragraph(paragraph, values)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def extract_xlsx_fields(raw: bytes) -> list:
    '''Список плейсхолдеров в бланке Excel — по всем листам.'''
    wb = load_workbook(io.BytesIO(raw))
    fields, seen = [], set()
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str) or '{{' not in cell.value:
                    continue
                for name in PLACEHOLDER_RE.findall(cell.value):
                    key = name.strip()
                    if key and key not in seen:
                        seen.add(key)
                        fields.append(key)
    return fields


def fill_xlsx(raw: bytes, values: dict) -> bytes:
    '''Заполняет бланк Excel. Меняем ТОЛЬКО значение ячейки — стиль, ширина колонок, объединения
    и формулы соседних ячеек остаются нетронутыми.'''
    wb = load_workbook(io.BytesIO(raw))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str) or '{{' not in cell.value:
                    continue

                def substitute(match):
                    name = match.group(1).strip()
                    return str(values[name]) if name in values else match.group(0)

                new_value = PLACEHOLDER_RE.sub(substitute, cell.value)
                if new_value == cell.value:
                    continue
                # Если в ячейке был ТОЛЬКО плейсхолдер и подставилось число — пишем именно число,
                # иначе Excel получит текст и не посчитает по нему сумму.
                cleaned = new_value.replace(' ', '').replace('\u00a0', '').replace(',', '.')
                try:
                    cell.value = float(cleaned) if '.' in cleaned else int(cleaned)
                except (TypeError, ValueError):
                    cell.value = new_value
    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()